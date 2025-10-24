"""Utilities for exporting meetings into various formats."""
from __future__ import annotations

from datetime import datetime, date, time, timedelta
from pathlib import Path
from textwrap import wrap
from typing import Iterable

from models import Meeting, Project, Track, Item, Section


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _parse_time(value: str | None, default: str) -> time:
    try:
        return datetime.strptime(value or default, "%H:%M").time()
    except Exception:
        return datetime.strptime(default, "%H:%M").time()


def _parse_date(value: str | None, default: date | None = None) -> date:
    default = default or date.today()
    try:
        y, m, d = [int(x) for x in (value or "").split("-")]
        return date(y, m, d)
    except Exception:
        return default


def _iter_section_items(meeting: Meeting) -> Iterable[tuple[Section, list[Item]]]:
    sections = sorted(meeting.sections, key=lambda s: s.order)
    items_by_section: dict[str, list[Item]] = {}
    for item in meeting.items:
        items_by_section.setdefault(item.section_name, []).append(item)
    for section in sections:
        items = sorted(items_by_section.get(section.name, []), key=lambda i: i.order)
        yield section, items


def _note_lines(item: Item) -> list[str]:
    lines: list[str] = []
    for note in item.notes:
        prefix = "Addendum" if note.is_addendum else f"Note ({note.meeting_date})"
        lines.append(f"{prefix}: {note.text.strip()}")
    return lines


def _escape_ics(value: str) -> str:
    return value.replace("\\", "\\\\").replace(";", "\\;").replace(",", "\\,")


def build_agenda_lines(project: Project, track: Track, meeting: Meeting) -> list[str]:
    lines = [
        f"Project: {project.name}",
        f"Track: {track.name}",
        f"Meeting: {meeting.number} — {meeting.header.date}",
        f"Time: {meeting.header.start} - {meeting.header.end}",
        f"Location: {meeting.header.location}",
    ]
    if meeting.header.teams_link:
        lines.append(f"Teams: {meeting.header.teams_link}")
    lines.append("")
    for section, items in _iter_section_items(meeting):
        lines.append(section.name or "Section")
        if not items:
            lines.append("  (No items)")
        for item in items:
            meta_bits = []
            if item.assignee_id:
                meta_bits.append(f"Assigned to {item.assignee_id}")
            if item.due_date:
                meta_bits.append(f"Due {item.due_date}")
            if item.tags:
                meta_bits.append(", ".join(filter(None, item.tags)))
            meta = f" ({'; '.join(meta_bits)})" if meta_bits else ""
            lines.append(f"  [{item.status}] {item.description}{meta}")
            for note_line in _note_lines(item):
                lines.append(f"    • {note_line}")
        lines.append("")
    return lines


def export_meeting_to_pdf(project: Project, track: Track, meeting: Meeting, path: Path) -> Path:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("ReportLab is required for PDF export. Install with 'pip install reportlab'.") from exc

    _ensure_parent(path)
    agenda_lines = build_agenda_lines(project, track, meeting)

    page_width, page_height = letter
    margin = 54
    line_height = 14

    c = canvas.Canvas(str(path), pagesize=letter)
    y = page_height - margin

    def draw_wrapped(text: str, font: str = "Helvetica", size: int = 11, bold: bool = False) -> None:
        nonlocal y
        if bold:
            font_name = "Helvetica-Bold"
        else:
            font_name = font
        c.setFont(font_name, size)
        max_width = page_width - (margin * 2)
        wrap_width = max(20, int(max_width / (size * 0.5)))
        for line in wrap(text, wrap_width) or [""]:
            if y < margin:
                c.showPage()
                c.setFont(font_name, size)
                y = page_height - margin
            c.drawString(margin, y, line)
            y -= line_height
        if not text:
            y -= 4

    draw_wrapped(meeting.header.topic or "Meeting Agenda", size=16, bold=True)
    y -= 6
    for line in agenda_lines:
        if not line:
            y -= 6
            continue
        is_section = line and not line.startswith("  ") and not line.startswith("Teams:") and not line.startswith("Project:") and not line.startswith("Track:") and not line.startswith("Meeting:") and not line.startswith("Time:") and not line.startswith("Location:")
        draw_wrapped(line, size=12 if is_section else 11, bold=is_section)
    c.save()
    return path


def export_meeting_to_docx(project: Project, track: Track, meeting: Meeting, path: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("python-docx is required for DOCX export. Install with 'pip install python-docx'.") from exc

    _ensure_parent(path)
    doc = Document()

    doc.add_heading(meeting.header.topic or "Meeting Agenda", level=0)
    doc.add_paragraph(f"Project: {project.name}")
    doc.add_paragraph(f"Track: {track.name}")
    doc.add_paragraph(f"Meeting: {meeting.number} — {meeting.header.date}")
    doc.add_paragraph(f"Time: {meeting.header.start} - {meeting.header.end}")
    doc.add_paragraph(f"Location: {meeting.header.location}")
    if meeting.header.teams_link:
        doc.add_paragraph(f"Teams link: {meeting.header.teams_link}")

    for section, items in _iter_section_items(meeting):
        doc.add_heading(section.name or "Section", level=2)
        if not items:
            doc.add_paragraph("(No items)")
            continue
        for item in items:
            details: list[str] = []
            if item.assignee_id:
                details.append(f"Assigned to {item.assignee_id}")
            if item.due_date:
                details.append(f"Due {item.due_date}")
            if item.tags:
                details.append(", ".join(filter(None, item.tags)))
            detail_text = f" ({'; '.join(details)})" if details else ""
            doc.add_paragraph(f"[{item.status}] {item.description}{detail_text}", style="List Bullet")
            for note_line in _note_lines(item):
                para = doc.add_paragraph(note_line, style="List Number")
                para.paragraph_format.left_indent = para.paragraph_format.left_indent or 0

    doc.save(str(path))
    return path


def export_meeting_to_ics(project: Project, track: Track, meeting: Meeting, path: Path) -> Path:
    _ensure_parent(path)

    meeting_date = _parse_date(meeting.header.date)
    start_time = _parse_time(meeting.header.start, "09:00")
    end_time = _parse_time(meeting.header.end, "10:00")
    start_dt = datetime.combine(meeting_date, start_time)
    end_dt = datetime.combine(meeting_date, end_time)
    if end_dt <= start_dt:
        end_dt = start_dt + timedelta(hours=1)

    description_lines = build_agenda_lines(project, track, meeting)
    description = "\\n".join(_escape_ics(line) for line in description_lines)
    summary = _escape_ics(meeting.header.topic or "Meeting")
    location = _escape_ics(meeting.header.location or "")

    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Meeting Manager (Projects)//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
        "BEGIN:VEVENT",
        f"UID:{meeting.id}@meetingmanager.local",
        f"DTSTAMP:{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}",
        f"DTSTART:{start_dt.strftime('%Y%m%dT%H%M%S')}",
        f"DTEND:{end_dt.strftime('%Y%m%dT%H%M%S')}",
        f"SUMMARY:{summary}",
        f"LOCATION:{location}",
        f"DESCRIPTION:{description}",
        "END:VEVENT",
        "END:VCALENDAR",
    ]

    path.write_text("\r\n".join(lines), encoding="utf-8")
    return path
